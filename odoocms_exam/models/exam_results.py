
from odoo import models, fields, api
import pdb

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Row, CT_Tc
from docx.shared import Inches, Pt, Cm

import logging
import os
import tempfile
import subprocess
import sys
import json
import requests

_logger = logging.getLogger(__name__)

endpoint = 'http://198.38.85.147:5000/image'
endpoint2 = 'http://198.38.85.147:5000/files/transcript.pdf'


def preventDocumentBreak(document):
    tags = document.element.xpath('//w:tr')
    rows = len(tags)
    for row in range(0, rows):
        tag = tags[row]  # Specify which <w:r> tag you want
        child = OxmlElement('w:cantSplit')  # Create arbitrary tag
        tag.append(child)  # Append in the new tag


def set_table_width(table, data):
    i = 0
    for w in data:
        set_column_width(table.columns[i], Cm(data[i]))
        i += 1


def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def shade_cells(cells, shade):
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement("w:shd")
        tcVAlign.set(qn("w:fill"), shade)
        tcPr.append(tcVAlign)


def row_border(row, t=True, l=True, b=True, r=True):
    cells = row._tr.getchildren()
    for cell in cells:
        if isinstance(cell, CT_Tc):
            tcPr = cell.tcPr
            tcBorders = OxmlElement('w:tcBorders')
            
            top = OxmlElement('w:top')
            if t:
                top.set(qn('w:val'), 'single')
            else:
                top.set(qn('w:val'), 'nil')
            
            left = OxmlElement('w:left')
            if l:
                left.set(qn('w:val'), 'single')
            else:
                left.set(qn('w:val'), 'nil')
            
            bottom = OxmlElement('w:bottom')
            if b:
                bottom.set(qn('w:val'), 'single')
            else:
                bottom.set(qn('w:val'), 'nil')
            
            right = OxmlElement('w:right')
            if r:
                right.set(qn('w:val'), 'single')
            else:
                right.set(qn('w:val'), 'nil')
            
            tcBorders.append(top)
            tcBorders.append(left)
            tcBorders.append(bottom)
            tcBorders.append(right)
            tcPr.append(tcBorders)


def transcript_border(table):
    tbl = table._tbl
    rows = tbl.getchildren()
    lastRow = len(table.rows) - 1
    
    for row in rows:
        if isinstance(row, CT_Row):
            cells = row.getchildren()
            
            if row.tr_idx == 0:  # Top Row
                for cell in cells:
                    if isinstance(cell, CT_Tc):
                        tcPr = cell.tcPr
                        tcBorders = OxmlElement('w:tcBorders')
                        
                        left = OxmlElement('w:left')
                        left.set(qn('w:val'), 'nil')
                        
                        right = OxmlElement('w:right')
                        right.set(qn('w:val'), 'nil')
                        
                        top = OxmlElement('w:top')
                        top.set(qn('w:val'), 'nil')
                        
                        tcBorders.append(left)
                        tcBorders.append(right)
                        tcBorders.append(top)
                        tcPr.append(tcBorders)
            
            elif row.tr_idx == lastRow:  # Bottom Row
                for cell in cells:
                    if isinstance(cell, CT_Tc):
                        tcPr = cell.tcPr
                        tcBorders = OxmlElement('w:tcBorders')
                        
                        top = OxmlElement('w:top')
                        top.set(qn('w:val'), 'nil')
                        
                        tcBorders.append(top)
                        tcPr.append(tcBorders)
            else:
                for cell in cells:
                    if isinstance(cell, CT_Tc):
                        tcPr = cell.tcPr
                        tcBorders = OxmlElement('w:tcBorders')
                        
                        top = OxmlElement('w:top')
                        top.set(qn('w:val'), 'nil')
                        bottom = OxmlElement('w:bottom')
                        bottom.set(qn('w:val'), 'nil')
                        
                        tcBorders.append(top)
                        tcBorders.append(bottom)
                        tcPr.append(tcBorders)


def add_row(table, data, fontSize=8):
    row_cells = table.add_row().cells
    i = 0
    for val in data:
        tb_cell_run = row_cells[i].paragraphs[0].add_run()
        tb_cell_run.add_text(str(str(data[i])))
        tb_cell_run.font.size = Pt(fontSize)
        i += 1


class OdooCMSExamGrade(models.Model):
    _name = 'odoocms.exam.grade'
    _description = 'Exam Grades'
    
    name = fields.Char('Grade Name')
    low_per = fields.Float('Per. From')
    high_per = fields.Float('Per. Below')
    gpa = fields.Float('GPA')


#Quiz=10, Assignments=15, Mid=25, Final=50  (Master)
class OdooCMSExamActivityType(models.Model):
    _name = 'odoocms.exam.activity.type'
    _description = 'Exam Activity Type'
    
    name = fields.Char('Activity Name', required=True)
    code = fields.Char('Activity Code', required=True)
    weightage = fields.Float('Weightage (%)',required=True)
    type = fields.Selection([('lecture','Lecture'),('lab','Lab')],'Type',default='lecture',required=True)
    
  
#Quiz=10, Assignments=15, Mid=25, Final=50  (Class Level)
class OdooCMSExamActivityClass(models.Model):
    _name = 'odoocms.exam.activity.class'
    _description = 'Exam Activity Class'
    _rec_name = 'activity_type_id'
    
    class_id = fields.Many2one('odoocms.class','Class')
    subject_id = fields.Many2one('odoocms.study.scheme.line','Subject',related='class_id.study_scheme_line_id',store=True)
    activity_type_id = fields.Many2one('odoocms.exam.activity.type','Activity Type')
    weightage = fields.Float('Weightage (%)')


class OdooCMSClass(models.Model):
    _inherit = 'odoocms.class'
    
    @api.multi
    @api.depends('exam_activity_class_ids','exam_activity_class_ids.weightage')
    def _get_planned_activities_sum(self):
        for rec in self:
            total = 0
            for activity in rec.exam_activity_class_ids:
                total += activity.weightage
            rec.activities_sum = total
            
    #Activities Planned
    exam_activity_class_ids = fields.One2many('odoocms.exam.activity.class', 'class_id', 'Activity Types')
    activities_sum = fields.Integer('Sum of Activities',compute='_get_planned_activities_sum')
    
    exam_activity_ids = fields.One2many('odoocms.exam.activity', 'class_id', 'Exam Activities')
    exam_activity_summary_ids = fields.One2many('odoocms.exam.activity.summary', 'class_id', 'Activity Summary')
    max_marks = fields.Float('Max Marks')

    def compute_result(self):
        max_marks = 0
        for student in self.student_ids:
            if student.total_marks > max_marks:
                max_marks = student.total_marks
        self.max_marks = max_marks

        for student in self.student_ids:
            student.normalized_marks = student.total_marks / max_marks * 100.0

            grade_rec = self.env['odoocms.exam.grade'].search([
                ('low_per','<=',student.normalized_marks),('high_per','>=',student.normalized_marks)])
            student.grade = grade_rec.name
            student.gpa = grade_rec.gpa
            # student.student_semester_id._get_sgpa

        self.convert_to_current()

    def finalize_result(self):
        self.compute_result()
        self.state = 'done'
        
    def generate_exam_activities(self):
        if not self.exam_activity_class_ids:
            for activity in self.env['odoocms.exam.activity.type'].search([]):
                data = {
                    'class_id': self.id,
                    'activity_type_id': activity.id,
                    'weightage': activity.weightage,
                }
                self.env['odoocms.exam.activity.class'].create(data)
            self.convert_to_current()
                
    @api.depends('exam_activity_class_ids')
    def convert_to_current(self):
        for rec in self:
            if rec.exam_activity_class_ids and rec.state == 'draft':
                rec.state = 'current'
    
    
class OdooCMSExamActivity(models.Model):
    _name = 'odoocms.exam.activity'
    _description = 'Exam Activity'
    
    name = fields.Char('Name')
    code = fields.Char('Code')
    activity_class_id = fields.Many2one('odoocms.exam.activity.class','Activity Class')
    
    class_id = fields.Many2one('odoocms.class','Class')
    academic_semester_id = fields.Many2one('odoocms.academic.semester', 'Academic Term', related='class_id.academic_semester_id',store=True)
    subject_id = fields.Many2one('odoocms.study.scheme.line','Subject',related='class_id.study_scheme_line_id',store=True)
    teacher_id = fields.Many2one('odoocms.faculty.staff','Teacher',related='class_id.faculty_staff_id',store=True)
    
    date_activity = fields.Date('Activity Date')
    max_marks = fields.Float('Max Marks')
    activity_lines = fields.One2many('odoocms.exam.activity.line','activity_id','Activity Lines')

    @api.multi
    def name_get(self):
        res = []
        for record in self:
            name = (record.name or '/') + ' - ' + (record.code or '')
            res.append((record.id, name))
        return res
    

class OdooCMSExamActivityLine(models.Model):
    _name = 'odoocms.exam.activity.line'
    _description = 'Exam Activity Line'

    activity_id = fields.Many2one('odoocms.exam.activity', 'Activity')
    student_id = fields.Many2one('odoocms.student','Student')
    summary_id = fields.Many2one('odoocms.exam.activity.summary','Activity Summary')
    activity_class_id = fields.Many2one('odoocms.exam.activity.class','Activity Type',
            related='activity_id.activity_class_id',store=True)
    max_marks = fields.Float('Max Marks',related='activity_id.max_marks',store=True)
    obtained_marks = fields.Float('Obtained Marks')
    percentage = fields.Float('Percentage',compute='_get_percentage',store=True)

    @api.depends('max_marks','obtained_marks')
    def _get_percentage(self):
        for rec in self:
            rec.percentage = (rec.obtained_marks or 0) / (rec.max_marks or 1) * 100
    
    
class OdooCMSExamActivitySummary(models.Model):
    _name = 'odoocms.exam.activity.summary'
    _description = 'CMS Exam Activity Summary'
    _rec_name = 'activity_class_id'
    
    class_id = fields.Many2one('odoocms.class', string='Class')
    subject_id = fields.Many2one('odoocms.study.scheme.line','Subject',related='class_id.study_scheme_line_id',store=True)
    academic_semester_id = fields.Many2one('odoocms.academic.semester', string='Academic Term',
                                           related='class_id.academic_semester_id', store=True)
    
    student_id = fields.Many2one('odoocms.student', string='Student')
    activity_class_id = fields.Many2one('odoocms.exam.activity.class','Activity Type')
    registration_id = fields.Many2one('odoocms.student.subject','Student Subject',compute='_get_student_subject',store=True)
    
    percentage = fields.Float('Percentage', compute='_get_percentage', store=True)
    activity_lines = fields.One2many('odoocms.exam.activity.line','summary_id','Activity Lines',)
    company_id = fields.Many2one('res.company', string='Company', default=lambda self: self.env['res.company']._company_default_get())

    @api.depends('class_id', 'student_id')
    def _get_student_subject(self):
        for result in self:
            if result.student_id and result.class_id:
                class_id = self.env['odoocms.student.subject'].search([
                    ('student_id','=',result.student_id.id),('class_id','=',result.class_id.id)])
                result.registration_id = class_id.id
                
    @api.depends('activity_lines', 'activity_lines.percentage')
    def _get_percentage(self):
        for result in self:
            percentage = cnt = 0
            for activity in result.activity_lines:
                percentage += activity.percentage
                cnt += 1
                
            result.percentage = percentage / (cnt or 1)
            

class OdooCMSStudentSubject(models.Model):
    _inherit = "odoocms.student.subject"

    exam_activity_ids = fields.One2many('odoocms.exam.activity.summary', 'registration_id','Exam Activities')
    total_marks = fields.Float('Total Marks',compute='_get_total_marks',store=True)
    normalized_marks = fields.Float('Normalized Marks')
    grade = fields.Char('Grade',size=5)
    grade_points = fields.Float('Grade Points',compute='_get_grade_points',store=True)
    gpa = fields.Float('GPA', digits=(8, 2))
    include_in_cgpa = fields.Boolean('Include in CGPA', default=True)
    earn_credit = fields.Boolean('Earn Credit', default=True)
    repeat_candidate = fields.Boolean('Repeat Candidate',default=False)
    valid_attempt = fields.Boolean('Valid Attempt',default=True)

    @api.depends('gpa', 'credits')
    def _get_grade_points(self):
        for rec in self:
            rec.grade_points = rec.gpa * rec.credits
        
    @api.depends('exam_activity_ids','exam_activity_ids.percentage')
    def _get_total_marks(self):
        for reg in self:
            total = 0
            for rec in reg.exam_activity_ids:
                total += rec.percentage * rec.activity_class_id.weightage
            reg.total_marks = total / 100.0


class OdooCMSStudentSemester(models.Model):
    _inherit = "odoocms.student.semester"

    grade_points = fields.Float('Grade Points',compute='_get_sgpa',store=True)
    credits = fields.Float('Semester Credits',compute='_get_sgpa',store=True)
    sgpa = fields.Float('SGPA',compute='_get_sgpa',store=True)

    cgp = fields.Float('CGP', compute='_get_sgpa', store=True)
    cch = fields.Float('CCH', compute='_get_sgpa', store=True)
    cgpa = fields.Float('CGPA', compute='_get_sgpa', store=True)
    
    @api.depends('student_subject_ids.grade_points','student_subject_ids.credits','student_subject_ids.state')
    def _get_sgpa(self):
        for rec in self:
            credits = grade_points = 0
            for subject in rec.student_subject_ids:
                credits += subject.credits
                grade_points += subject.grade_points
                
            rec.credits = credits
            rec.grade_points = grade_points
            rec.sgpa = grade_points / (credits or 1.00)
            
            cch = cgp = 0
            for semester in rec.student_id.semester_ids.filtered(
                    lambda l: l.semester_id.number <= rec.semester_id.number):
                cch += semester.credits
                cgp += semester.grade_points

            rec.cch = cch
            rec.cgp = cgp
            rec.cgpa = cgp / (cch or 1.00)
        
    
class OdooCMSStudent(models.Model):
    _inherit = 'odoocms.student'

    grade_points = fields.Float('Grade Points', compute='_get_cgpa', store=True)
    credits = fields.Float('Earned Credits', compute='_get_cgpa', store=True)
    cgpa = fields.Float('CGPA', compute='_get_cgpa', store=True)

    @api.depends('semester_ids','semester_ids.grade_points', 'semester_ids.credits')
    def _get_cgpa(self):
        for rec in self:
            credits = grade_points = 0
            for semester in rec.semester_ids:
                credits += semester.credits
                grade_points += semester.grade_points
        
            rec.credits = credits
            rec.grade_points = grade_points
            rec.cgpa = grade_points / (credits or 1.00)
            
    def prereq_satisfy(self, prereq):
        return not prereq in self.result_subject_ids.filtered(lambda l: l.include_in_cgpa and l.grade in ('W','F')).mapped('subject_id')

    def doc5(self):
        document = Document()
        section = document.sections[-1]
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    
        obj_styles = document.styles
        obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(8)
        obj_font.name = 'Times New Roman'
    
        obj_charstyle = obj_styles.add_style('CommentsStyle2', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(10)
        obj_font.name = 'Times New Roman'
    
        header = document.sections[0].header
        paragraph = header.add_paragraph()
        paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        paragraph.add_run(self.env.user.company_id.name).bold = True
    
        paragraph = header.add_paragraph()
        paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
        paragraph.add_run('Registration No: ')
        paragraph.add_run(self.id_number or ' ').bold = True
    
        # paragraph = header.add_paragraph('Transcript', style='Intense Quote')
        # paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    
        header_table = header.add_table(1, 3, width=500)
        header_table.autofit = False
    
        A = header_table.cell(0, 0)
        pt = A.paragraphs[0]
        t = pt.text = ''
        pt.add_run("Student's Name: ", style='CommentsStyle').bold = True
        pt.add_run(self.name, style='CommentsStyle').bold = False
        pt.paragraph_format.space_before = Pt(0)
        pt.paragraph_format.space_after = Pt(0)
    
        paragraph = A.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run("Program: ", style='CommentsStyle').bold = True
        paragraph.add_run(self.program_id.name, style='CommentsStyle').bold = False
    
        paragraph = A.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run("Plan: ", style='CommentsStyle').bold = True
        paragraph.add_run(self.program_id.name + ' Major', style='CommentsStyle').bold = False
    
        B = header_table.cell(0, 2)
        pt = B.paragraphs[0]
        t = pt.text = ''
        pt.add_run("Father's Name: ", style='CommentsStyle').bold = True
        pt.add_run(self.father_name, style='CommentsStyle').bold = False
        pt.paragraph_format.space_before = Pt(0)
        pt.paragraph_format.space_after = Pt(0)
    
        cells = header_table.add_row().cells
        cells[0]._element.clear_content()
        table = cells[0].add_table(rows=0, cols=4)
        table.style = 'Table Grid'
        table.autofit = False
    
        set_table_width(table, [1.8, 5.3, 1.0, 1.0])
        add_row(table, ['Code', 'Title', 'CH', 'Grd'])
    
        cells[2]._element.clear_content()
        table = cells[2].add_table(rows=0, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        table.style = 'Table Grid'
        table.autofit = False
    
        set_table_width(table, [1.8, 5.3, 1.0, 1.0])
        add_row(table, ['Code', 'Title', 'CH', 'Grd'])
    
        set_table_width(header_table, [9.4, 1.0, 9.4])
    
        footer = document.sections[0].footer
        footer.is_linked_to_previous = False
    
        pt = footer.paragraphs[0]
        pt.add_run('"The Official Transcript carries the embossed stamp of the University"', style='CommentsStyle').bold = True
        pt.paragraph_format.space_before = Pt(0)
        pt.paragraph_format.space_after = Pt(0)
    
        paragraph = footer.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('Transcript Prepared By: ---------------------------------------------', style='CommentsStyle').bold = False
    
        paragraph = footer.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('Transcript Checked By: ---------------------------------------------', style='CommentsStyle').bold = False
    
        paragraph = footer.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('Date of Issue: '+ str(fields.Date.today()), style='CommentsStyle').bold = False
    
        paragraph = footer.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('"Errors and Omissions are subject to Subsequent rectification"', style='CommentsStyle').bold = True
        paragraph.add_run("\t\t\tController of Examinations", style='CommentsStyle2').bold = True
    
        big_table = document.add_table(0, 1)
        big_table.autofit = False
        set_table_width(big_table, [9.5])
    
        for semester in self.semester_ids:
            row = big_table.add_row()
        
            tag = row._tr
            child = OxmlElement('w:cantSplit')  # Create arbitrary tag
            tag.append(child)
        
            cells = row.cells
            cells[0]._element.clear_content()
        
            # label = cells[0].add_paragraph()
            # label.paragraph_format.keep_with_next = True
            # label.paragraph_format.space_before = Pt(0)
            # label.paragraph_format.space_after = Pt(0)
        
            table = cells[0].add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            a = table.cell(0, 0)
            b = table.cell(0, 3)
            A = a.merge(b)
            A.text = semester.academic_semester_id.name
            A.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        
            for subject in semester.student_subject_ids:
                add_row(table, [
                    subject.subject_id.subject_id.code,
                    subject.subject_id.subject_id.name,
                    subject.subject_id.weightage,
                    subject.grade
                ])
        
            set_table_width(table, [1.8, 5.5, 1.0, 1.0])
            transcript_border(table)

            row = table.add_row()
            row_border(row, b=False)
            a = table.cell(len(table.rows) - 1, 0)
            b = table.cell(len(table.rows) - 1, 3)
            A = a.merge(b)

            tb_cell_run = A.paragraphs[0].add_run()
            tb_cell_run.add_text("\tSCH: " + str(semester.credits))
            tb_cell_run.add_text("\t\tSGP: " + '{0:,.2f}'.format(semester.grade_points))
            tb_cell_run.add_text("\tSGPA: " + '{0:,.2f}'.format(semester.sgpa))
            tb_cell_run.font.size = Pt(8)

            

            row = table.add_row()
            # row_border(table.rows[len(table.rows)-1],t=False)
            row_border(row, t=False)
            a = table.cell(len(table.rows) - 1, 0)
            b = table.cell(len(table.rows) - 1, 3)
            A = a.merge(b)

            tb_cell_run = A.paragraphs[0].add_run()
            tb_cell_run.add_text("\tCCH: " + str(semester.cch))
            tb_cell_run.add_text("\t\tCGP: " + '{0:,.2f}'.format(semester.cgp))
            tb_cell_run.add_text("\tCGPA: " + '{0:,.2f}'.format(semester.cgpa))
            tb_cell_run.font.size = Pt(8)
        
            for row in table.rows:
                row.height = Cm(0.4)
        
            label = cells[0].paragraphs[0]
            label.paragraph_format.keep_with_next = True
            label.paragraph_format.space_before = Pt(0)
            label.paragraph_format.space_after = Pt(0)
    
        sectPr = document.sections[-1]._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')

        preventDocumentBreak(document)

        # document.save('demo.docx')
    
        temporary_files = []
    
        doc_report_fd, doc_report_path = tempfile.mkstemp(suffix='.docx', prefix='report.tmp.')
        os.close(doc_report_fd)
        temporary_files.append(doc_report_path)

        pdf_report_fd, pdf_report_path = tempfile.mkstemp(suffix='.pdf', prefix='report.tmp.')
        os.close(pdf_report_fd)
        temporary_files.append(pdf_report_path)
        
    
        document.save(doc_report_path)
    
        #send to server
        headers = {
            'Content-Type': 'multipart/form-data',
        }

        response = requests.put(endpoint, files={'file': open(doc_report_path, 'rb')})
        if response.status_code == 200:
            print(response.text)

        response = requests.get(endpoint2)
        print(response.status_code)

        if response.status_code == 200:
            with open(pdf_report_path, 'wb') as out_file:  # change file name for PNG images
                out_file.write(response.content)
 
       
        #
        # try:
        #     # wkhtmltopdf = [_get_wkhtmltopdf_bin()] + command_args + files_command_args + paths + [pdf_report_path]
        #     wkhtmltopdf = ["/usr/bin/unoconv", "-f", "pdf", "-o", pdf_report_path, doc_report_path]
        #     process = subprocess.Popen(wkhtmltopdf, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        #     out, err = process.communicate()
        # except:
        #     raise

        with open(pdf_report_path, 'rb') as pdf_document:
            pdf_content = pdf_document.read()
    
        # Manual cleanup of the temporary files
        for temporary_file in temporary_files:
            try:
                os.unlink(temporary_file)
            except (OSError, IOError):
                _logger.error('Error when trying to remove file %s' % temporary_file)

        return pdf_content
    
